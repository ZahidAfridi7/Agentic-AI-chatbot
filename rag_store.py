# rag_store.py
import os
from typing import List, Optional, Tuple
from io import BytesIO
from langchain_core.documents import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings
STORAGE_DIR = os.getenv("RAG_STORAGE_DIR", "./storage")

# --- Embeddings (single source of truth) ---
def get_embeddings():
    # Free, fast, widely used
    return HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")

def _dataset_path(dataset_id: str) -> str:
    return os.path.join(STORAGE_DIR, dataset_id)

def get_vectorstore(dataset_id: str):
    path = _dataset_path(dataset_id)
    if not os.path.isdir(path):
        os.makedirs(path, exist_ok=True)
        return None
    try:
        # IMPORTANT: provide the same embedding function when loading,
        # so add/search keep working after reload.
        return FAISS.load_local(
            path,
            get_embeddings(),
            allow_dangerous_deserialization=True
        )
    except Exception:
        return None

def save_vectorstore(dataset_id: str, vs: FAISS):
    path = _dataset_path(dataset_id)
    os.makedirs(path, exist_ok=True)
    vs.save_local(path)

def upsert_texts(
    dataset_id: str,
    texts: List[str],
    metadatas: Optional[List[dict]] = None,
    chunk_size: int = 1200,
    chunk_overlap: int = 150
):
    splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    docs: List[Document] = []
    for i, t in enumerate(texts):
        chunks = splitter.split_text((t or "").strip())
        meta = (metadatas[i] if metadatas and i < len(metadatas) else {}) or {}
        docs.extend([Document(page_content=c, metadata=meta) for c in chunks])

    embedder = get_embeddings()
    vs = get_vectorstore(dataset_id)
    if vs is None:
        vs = FAISS.from_documents(docs, embedder)
    else:
        # Ensure the vectorstore has an embedding_function attached
        if getattr(vs, "embedding_function", None) is None:
            vs.embedding_function = embedder
        vs.add_documents(docs)

    save_vectorstore(dataset_id, vs)
    return {"chunks_added": len(docs)}

def clear_dataset(dataset_id: str):
    path = _dataset_path(dataset_id)
    if os.path.isdir(path):
        for root, dirs, files in os.walk(path, topdown=False):
            for f in files:
                os.remove(os.path.join(root, f))
            for d in dirs:
                os.rmdir(os.path.join(root, d))
        os.rmdir(path)
    return {"cleared": True}

def get_retriever(dataset_id: str):
    vs = get_vectorstore(dataset_id)
    if vs is None:
        return None
    # k can be tuned later
    return vs.as_retriever(search_kwargs={"k": 4})

# --- File text extraction helpers ---

ALLOWED_EXTS = {".txt", ".docx", ".pdf"}

def _ext_ok(filename: str) -> bool:
    _, ext = os.path.splitext(filename.lower())
    return ext in ALLOWED_EXTS

def extract_plain_text(file_bytes: bytes, filename: str) -> Tuple[str, str]:
    """
    Return (text, ext). Raises ValueError for unsupported or empty content.
    """
    name = filename or "uploaded"
    _, ext = os.path.splitext(name.lower())

    if ext == ".txt":
        try:
            text = file_bytes.decode("utf-8", errors="ignore")
        except Exception:
            text = file_bytes.decode("latin-1", errors="ignore")
        text = text.strip()
        if not text:
            raise ValueError(f"{name}: no text content found")
        return text, ext

    elif ext == ".docx":
        try:
            from docx import Document as DocxDocument
        except ImportError as e:
            raise RuntimeError("python-docx not installed. Run: pip install python-docx") from e
        doc = DocxDocument(BytesIO(file_bytes))
        text = "\n".join(p.text for p in doc.paragraphs).strip()
        if not text:
            raise ValueError(f"{name}: no text content found in DOCX")
        return text, ext

    elif ext == ".pdf":
        try:
            from pypdf import PdfReader
        except ImportError as e:
            raise RuntimeError("pypdf not installed. Run: pip install pypdf") from e
        reader = PdfReader(BytesIO(file_bytes))
        chunks = []
        for page in reader.pages:
            try:
                chunks.append(page.extract_text() or "")
            except Exception:
                chunks.append("")
        text = "\n".join(chunks).strip()
        if not text:
            raise ValueError(f"{name}: no extractable text (image-based PDF?)")
        return text, ext

    else:
        raise ValueError(f"Unsupported file type: {ext}. Allowed: {', '.join(sorted(ALLOWED_EXTS))}")
